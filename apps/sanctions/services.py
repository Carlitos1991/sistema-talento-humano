import os
from datetime import date
from types import SimpleNamespace
from pathlib import Path
from xml.sax.saxutils import escape

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.table import Table as DocxTable
from docx.text.paragraph import Paragraph as DocxParagraph
from django.conf import settings
from pypdf import PdfReader, PdfWriter

PLACEHOLDER_KEYS = [
    'EMPLOYEE_FULL_NAME',
    'EMPLOYEE_FIRST_NAME',
    'EMPLOYEE_LAST_NAME',
    'EMPLOYEE_DOCUMENT_NUMBER',
    'EMPLOYEE_POSITION',
    'EMPLOYEE_UNIT',
    'REGIME_CODE',
    'REGIME_NAME',
    'NOTIFICATION_NAME',
    'MONTH_NAME',
    'MONTH_NUMBER',
    'YEAR',
    'REGISTRATION_DATE',
    'AUTHORITY_1_NAME',
    'AUTHORITY_1_POSITION',
    'AUTHORITY_2_NAME',
    'AUTHORITY_2_POSITION',
    'MINUTES_LATE',
    'REGS_WITHOUT_MARK',
    'OBSERVATIONS',
]


def _format_spanish_date(date_value):
    months = {
        1: 'enero',
        2: 'febrero',
        3: 'marzo',
        4: 'abril',
        5: 'mayo',
        6: 'junio',
        7: 'julio',
        8: 'agosto',
        9: 'septiembre',
        10: 'octubre',
        11: 'noviembre',
        12: 'diciembre',
    }
    if not date_value:
        return ''
    return f"{date_value.day:02d} de {months.get(date_value.month, '')} de {date_value.year}"


def build_notification_replacements(data):
    sequence_code = str(data.get('sequence_code', '') or '')
    user_code = str(data.get('user_code', '') or '')
    registration_date = data.get('registration_date', '')
    minutes_late = data.get('minutes_late')
    regs_without_mark = data.get('regs_without_mark')

    minutes_late_text = '0' if minutes_late is None else str(minutes_late)
    regs_without_mark_text = '0' if regs_without_mark is None else str(regs_without_mark)

    replacements = {
        '[FULL_NAME]': data.get('employee_full_name', ''),
        '[NAME]': data.get('employee_first_name', ''),
        '[LAST_NAME]': data.get('employee_last_name', ''),
        '[DOCUMENT_NUMBER]': data.get('employee_document_number', ''),
        '[POSITION]': data.get('employee_position', ''),
        '[UNIT]': data.get('employee_unit', ''),
        '[EMPLOYEE_FULL_NAME]': data.get('employee_full_name', ''),
        '[EMPLOYEE_FIRST_NAME]': data.get('employee_first_name', ''),
        '[EMPLOYEE_LAST_NAME]': data.get('employee_last_name', ''),
        '[EMPLOYEE_DOCUMENT_NUMBER]': data.get('employee_document_number', ''),
        '[EMPLOYEE_POSITION]': data.get('employee_position', ''),
        '[EMPLOYEE_UNIT]': data.get('employee_unit', ''),
        '[REGIME_CODE]': data.get('regime_code', ''),
        '[REGIME_NAME]': data.get('regime_name', ''),
        '[NOTIFICATION_NAME]': data.get('notification_name', ''),
        '[MONTH_NAME]': data.get('month_name', ''),
        '[MONTH_NUMBER]': str(data.get('month_number', '') or ''),
        '[YEAR]': str(data.get('year', '') or ''),
        '[AÑO]': str(data.get('year', '') or ''),
        '[REGISTRATION_DATE]': registration_date,
        '[DATE]': registration_date,
        '[today]': registration_date,
        '[TODAY]': registration_date,
        '[LOCALIDAD]': data.get('location', 'Loja') or 'Loja',
        '[SECUENCIA]': sequence_code,
        '[SECUENCIAL]': sequence_code,
        '[CODIGO]': sequence_code,
        '[CODIGO_USUARIO]': user_code,
        '[AUTHORITY_1_NAME]': data.get('authority_1_name', ''),
        '[AUTHORITY_1_POSITION]': data.get('authority_1_position', ''),
        '[AUTHORITY_2_NAME]': data.get('authority_2_name', ''),
        '[AUTHORITY_2_POSITION]': data.get('authority_2_position', ''),
        '[NOMBRE_AUTORIDAD1]': data.get('authority_1_name', ''),
        '[CARGO_AUTORIDAD1]': data.get('authority_1_position', ''),
        '[NOMBRE_AUTORIDAD2]': data.get('authority_2_name', ''),
        '[CARGO_AUTORIDAD2]': data.get('authority_2_position', ''),
        '[MINUTES_LATE]': minutes_late_text,
        '[MINUTOS]': minutes_late_text,
        '[REGS_WITHOUT_MARK]': regs_without_mark_text,
        '[OBSERVATIONS]': data.get('observations', ''),
    }
    return replacements


def build_replacements_from_mappings(data, mappings):
    base = build_notification_replacements(data)
    dynamic_replacements = {}

    for mapping in mappings:
        source_value = data.get(mapping.source_key, '')
        dynamic_replacements[mapping.placeholder] = str(source_value or '')

    base.update(dynamic_replacements)
    return base


def _get_nested_value(data, path):
    current = data
    for part in path.split('.'):
        if current is None:
            return ''
        if isinstance(current, dict):
            current = current.get(part)
        else:
            current = getattr(current, part, '')
    return current if current is not None else ''


def _split_expression(expression):
    parts = []
    current = []
    quote = None

    for char in expression:
        if char in ('"', "'"):
            if quote is None:
                quote = char
            elif quote == char:
                quote = None
            current.append(char)
            continue

        if char == '+' and quote is None:
            token = ''.join(current).strip()
            if token:
                parts.append(token)
            current = []
            continue

        current.append(char)

    token = ''.join(current).strip()
    if token:
        parts.append(token)
    return parts


def evaluate_mapping_expression(expression, data):
    expression = (expression or '').strip()
    if not expression:
        return ''

    if expression.lower() == 'today':
        return _format_spanish_date(date.today())

    if expression.lower().startswith('today:'):
        date_format = expression.split(':', 1)[1].strip() or '%d/%m/%Y'
        return date.today().strftime(date_format)

    value_parts = []
    for token in _split_expression(expression):
        if token.startswith(('"', "'")) and token.endswith(('"', "'")) and len(token) >= 2:
            value_parts.append(token[1:-1])
            continue
        if token.lower() == 'today':
            value_parts.append(_format_spanish_date(date.today()))
            continue
        if token.lower().startswith('today:'):
            date_format = token.split(':', 1)[1].strip() or '%d/%m/%Y'
            value_parts.append(date.today().strftime(date_format))
            continue
        value_parts.append(str(_get_nested_value(data, token)).strip())

    return ''.join(value_parts)


def build_replacements_from_global_mappings(data, mappings):
    replacements = build_notification_replacements(data)
    for mapping in mappings:
        replacements[mapping.placeholder] = evaluate_mapping_expression(mapping.expression, data)
    return replacements


def _replace_in_text(text, replacements):
    if not text:
        return text
    for placeholder, value in replacements.items():
        text = text.replace(placeholder, str(value or ''))
    return text


def _replace_in_paragraph(paragraph, replacements):
    original_text = ''.join(run.text for run in paragraph.runs)
    replaced_text = _replace_in_text(original_text, replacements)

    if original_text == replaced_text:
        for run in paragraph.runs:
            run.text = _replace_in_text(run.text, replacements)
        return

    if not paragraph.runs:
        return

    paragraph.runs[0].text = replaced_text
    for run in paragraph.runs[1:]:
        run.text = ''


def _replace_in_table(table, replacements):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                _replace_in_paragraph(paragraph, replacements)
            for inner_table in cell.tables:
                _replace_in_table(inner_table, replacements)


def fill_docx_template(template_path, output_path, replacements):
    document = Document(template_path)

    for paragraph in document.paragraphs:
        _replace_in_paragraph(paragraph, replacements)

    for table in document.tables:
        _replace_in_table(table, replacements)

    for section in document.sections:
        for paragraph in section.header.paragraphs:
            _replace_in_paragraph(paragraph, replacements)
        for table in section.header.tables:
            _replace_in_table(table, replacements)
        for paragraph in section.footer.paragraphs:
            _replace_in_paragraph(paragraph, replacements)
        for table in section.footer.tables:
            _replace_in_table(table, replacements)

    document.save(output_path)


def build_dynamic_notification_docx(output_path, sections, replacements, sequence_code, regime_code, year, user_code,
                                    registration_date, location='Loja'):
    """
    Construye un DOCX dinámico sin plantilla Word subida por el usuario.
    Encabezado fijo:
    NOTIFICACIÓN Nº 0001-BIO-2026-JUPE
    Loja, 08/04/2026
    """
    document = Document()

    full_sequence = f'{sequence_code}-{regime_code}-{year}-{user_code}'
    replacements = {**replacements, '[CODIGO_NOTIFICACION]': full_sequence}

    # Encabezado principal
    header = document.add_paragraph(f'NOTIFICACIÓN Nº {full_sequence}')
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if header.runs:
        header.runs[0].bold = True
        header.runs[0].font.size = Pt(14)
    header.paragraph_format.space_after = Pt(12)

    # Localidad y fecha
    date_line = document.add_paragraph(f'{location}, {registration_date}')
    date_line.alignment = WD_ALIGN_PARAGRAPH.LEFT
    date_line.paragraph_format.space_after = Pt(14)

    # Secciones dinámicas configuradas en interfaz personalizada
    for section in sections:
        content = section.content or ''
        for placeholder, value in replacements.items():
            content = content.replace(placeholder, str(value or ''))

        paragraph = document.add_paragraph(content)
        if section.section_type == 'TITLE':
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            if paragraph.runs:
                paragraph.runs[0].bold = True
                paragraph.runs[0].font.size = Pt(12)
            paragraph.paragraph_format.space_before = Pt(12)
            paragraph.paragraph_format.space_after = Pt(6)
        else:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            paragraph.paragraph_format.space_after = Pt(8)

    document.save(output_path)


def extract_docx_preview(template_path, max_blocks=24):
    try:
        document = Document(template_path)
    except Exception:
        return []

    blocks = []

    for block in document.element.body.iterchildren():
        if len(blocks) >= max_blocks:
            break

        tag = block.tag.split('}')[-1]
        if tag == 'p':
            paragraph = DocxParagraph(block, document)
            text = paragraph.text.strip()
            if text:
                blocks.append({'type': 'paragraph', 'text': text})
        elif tag == 'tbl':
            table = DocxTable(block, document)
            rows = []
            for row in table.rows:
                rows.append([cell.text.strip() for cell in row.cells])
            if rows:
                blocks.append({'type': 'table', 'rows': rows})

    return blocks


